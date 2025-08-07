import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document

st.title("Registro de Plantões")

# ---------------------------
# Estado inicial
# ---------------------------
if "registros" not in st.session_state:
    st.session_state.registros = []

# ---------------------------
# Configurações iniciais
# ---------------------------
with st.expander("Configurações iniciais (preencher uma vez)", expanded=True):
    nome = st.text_input("Nome do médico", key="nome")
    hospital = st.text_input("Hospital", key="hospital", placeholder="Ex.: HOE")
    mes_ano = st.text_input("Mês/Ano do relatório (ex.: Julho/2025)", key="mes_ano")
    valor_hora = st.number_input("Valor da hora (R$)", key="valor_hora", min_value=0.0, step=10.0, value=160.0)
    produtividade = st.number_input("Produtividade do mês (R$)", key="produtividade", min_value=0.0, step=50.0)

st.markdown("---")

# ---------------------------
# Cadastro de plantões
# ---------------------------
st.subheader("Novo plantão")

col1, col2 = st.columns([1, 1])
with col1:
    data = st.date_input("Data do plantão")
with col2:
    local = st.selectbox("Local de trabalho", ["Ambulatório", "Centro Cirúrgico", "Diarismo"])

horario = st.selectbox(
    "Horário do plantão",
    ["07:00 - 13:00", "13:00 - 19:00", "07:00 - 19:00", "19:00 - 07:00"]
)

# calcula horas automaticamente a partir do horário escolhido
map_horas = {
    "07:00 - 13:00": 6,
    "13:00 - 19:00": 6,
    "07:00 - 19:00": 12,
    "19:00 - 07:00": 12,
}
horas_calculadas = map_horas[horario]
st.caption(f"Horas calculadas: **{horas_calculadas}h**")

if st.button("Registrar plantão"):
    st.session_state.registros.append({
        "Data": data,             # datetime.date
        "Horário": horario,       # string
        "Horas": horas_calculadas,
        "Local": local            # string
    })
    st.success("Plantão registrado.")

st.markdown("---")

# ---------------------------
# Lista de plantões + excluir
# ---------------------------
if st.session_state.registros:
    st.subheader("Plantões registrados")
    df = pd.DataFrame(st.session_state.registros)
    df_exibe = df.copy()
    df_exibe["Data"] = df_exibe["Data"].apply(lambda d: d.strftime("%d/%m/%Y"))
    df_exibe = df_exibe[["Data", "Horário", "Horas", "Local"]]

    # Mostra a lista com botões de excluir
    for i, row in df_exibe.iterrows():
        c1, c2, c3, c4, c5 = st.columns([1.2, 1.3, 0.6, 1.6, 0.8])
        c1.write(row["Data"])
        c2.write(row["Horário"])
        c3.write(f'{int(row["Horas"])}h')
        c4.write(row["Local"])
        if c5.button("Excluir", key=f"del_{i}"):
            st.session_state.registros.pop(i)
            st.rerun()

# ---------------------------
# Relatório no estilo solicitado
# ---------------------------
if st.session_state.registros:
    st.markdown("---")
    st.subheader("Relatório (visual)")

    df = pd.DataFrame(st.session_state.registros)
    df["Valor"] = df["Horas"] * float(valor_hora)

    # Ordem dos setores
    setores_ordem = ["Diarismo", "Ambulatório", "Centro Cirúrgico"]
    total_geral_valor = 0.0

    for setor in setores_ordem:
        df_setor = df[df["Local"] == setor].copy()
        if df_setor.empty:
            continue
        # ordenar por data
        df_setor = df_setor.sort_values(by="Data")
        st.markdown(f"### *{setor}*")
        total_horas_setor = int(df_setor["Horas"].sum())
        valor_setor = df_setor["Valor"].sum()
        total_geral_valor += valor_setor

        for _, r in df_setor.iterrows():
            st.write(f"{r['Data'].strftime('%d/%m/%Y')} ({r['Horário']}) - {int(r['Horas']):02d}h")

        st.write(f"**Total: {total_horas_setor} horas**")
        st.write(f"**Valor: {total_horas_setor} X {valor_hora:.0f} = R$ {valor_setor:,.2f}**")
        st.write("")

    st.markdown(f"**Valor total: R$ {total_geral_valor:,.2f}**")
    st.markdown(f"**Produtividade**  \n{mes_ano or 'Mês/Ano não informado'}: R$ {float(produtividade):,.2f}")
    st.markdown(f"**Valor final: R$ {(total_geral_valor + float(produtividade)):,.2f}**")

    # ---------------------------
    # Botão: gerar .docx p/ download
    # ---------------------------
    st.markdown("---")
    st.subheader("Gerar relatório Word (.docx)")

    def gerar_docx():
        doc = Document()
        # Cabeçalho
        cabecalho = f"Serviços {hospital or ''} - {nome or ''} - {mes_ano or ''}".strip()
        doc.add_heading(cabecalho if cabecalho else "Serviços", level=1)

        total_geral_val = 0.0
        for setor in setores_ordem:
            df_set = df[df["Local"] == setor].copy()
            if df_set.empty:
                continue
            df_set = df_set.sort_values(by="Data")
            doc.add_paragraph(f"*{setor}*")

            total_horas_set = int(df_set["Horas"].sum())
            valor_set = df_set["Valor"].sum()
            total_geral_val += valor_set

            for _, r in df_set.iterrows():
                linha = f"{r['Data'].strftime('%d/%m/%Y')} ({r['Horário']}) - {int(r['Horas']):02d}h"
                doc.add_paragraph(linha)

            doc.add_paragraph(f"Total: {total_horas_set} horas")
            doc.add_paragraph(f"Valor: {total_horas_set} X {int(valor_hora)} = {valor_set:,.2f}\n")

        doc.add_paragraph(f"Valor total: {total_geral_val:,.2f}")
        doc.add_paragraph(f"\nProdutividade\n{mes_ano or ''}: {float(produtividade):,.2f}")
        doc.add_paragraph(f"\nValor final: {total_geral_val + float(produtividade):,.2f}")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    if st.button("Gerar arquivo .docx"):
        buf = gerar_docx()
        nome_arquivo = f"Relatorio_{(nome or 'medico').replace(' ', '_')}.docx"
        st.download_button("Baixar relatório .docx", data=buf, file_name=nome_arquivo, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
else:
    st.info("Ainda não há plantões registrados.")'
